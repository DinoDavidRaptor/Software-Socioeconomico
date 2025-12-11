"""
Módulo de exportación a Word (DOCX).
Autor: DINOS Tech
Versión: 0.2.0
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict


class ExportadorWord:
    """Clase para exportar estudios socioeconómicos a formato Word."""
    
    def __init__(self, config_empresa: Dict):
        """
        Inicializa el exportador con la configuración de la empresa.
        
        Args:
            config_empresa: Diccionario con datos de la empresa.
        """
        self.config = config_empresa
    
    def _agregar_encabezado(self, doc: Document):
        """
        Agrega el encabezado con información de la empresa.
        
        Args:
            doc: Documento de Word.
        """
        # Logo
        logo_path = self.config.get("logo", "")
        if logo_path and os.path.exists(logo_path):
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(logo_path, width=Inches(1.5))
            except:
                pass
        
        # Nombre de la empresa
        heading = doc.add_heading(self.config.get("nombre", ""), level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Datos de contacto
        contact = doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_text = contact.add_run(
            f"{self.config.get('direccion', '')}\n"
            f"Tel: {self.config.get('telefono', '')} | Email: {self.config.get('email', '')}"
        )
        contact_text.font.size = Pt(10)
        
        doc.add_paragraph()
    
    def _agregar_seccion(self, doc: Document, titulo: str):
        """
        Agrega un título de sección al documento.
        
        Args:
            doc: Documento de Word.
            titulo: Título de la sección.
        """
        heading = doc.add_heading(titulo, level=2)
        heading_format = heading.runs[0].font
        heading_format.color.rgb = RGBColor(44, 62, 80)
    
    def _agregar_tabla_datos(self, doc: Document, datos: list):
        """
        Agrega una tabla de datos al documento.
        
        Args:
            doc: Documento de Word.
            datos: Lista de tuplas (etiqueta, valor).
        """
        table = doc.add_table(rows=len(datos), cols=2)
        table.style = 'Light Grid Accent 1'
        
        for idx, (etiqueta, valor) in enumerate(datos):
            row = table.rows[idx]
            cell_label = row.cells[0]
            cell_value = row.cells[1]
            
            # Etiqueta en negrita
            cell_label.text = str(etiqueta)
            for paragraph in cell_label.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
            
            # Valor
            cell_value.text = str(valor)
        
        doc.add_paragraph()
    
    def _agregar_datos_personales(self, doc: Document, datos: Dict):
        """Agrega la sección de datos personales."""
        self._agregar_seccion(doc, "DATOS PERSONALES")
        
        dp = datos.get("datos_personales", {})
        
        datos_tabla = [
            ("Nombre Completo:", dp.get("nombre_completo", "N/A")),
            ("Edad:", dp.get("edad", "N/A")),
            ("Fecha de Nacimiento:", dp.get("fecha_nacimiento", "N/A")),
            ("Género:", dp.get("genero", "N/A")),
            ("Estado Civil:", dp.get("estado_civil", "N/A")),
            ("Nacionalidad:", dp.get("nacionalidad", "N/A")),
            ("Lugar de Nacimiento:", dp.get("lugar_nacimiento", "N/A")),
            ("CURP:", dp.get("curp", "N/A")),
            ("RFC:", dp.get("rfc", "N/A")),
            ("INE:", dp.get("ine", "N/A")),
            ("NSS:", dp.get("nss", "N/A")),
            ("Teléfono:", dp.get("telefono", "N/A")),
            ("Email:", dp.get("email", "N/A")),
            ("Dirección:", dp.get("direccion", "N/A")),
            ("Escolaridad:", dp.get("escolaridad", "N/A")),
            ("Carrera/Especialidad:", dp.get("carrera_especialidad", "N/A")),
            ("Estado de Estudios:", dp.get("estado_estudios", "N/A")),
            ("Licencia de Conducir:", "Sí" if dp.get("licencia_conducir") else "No")
        ]
        
        if dp.get("licencia_conducir"):
            datos_tabla.extend([
                ("Tipo Licencia:", dp.get("licencia_tipo", "N/A")),
                ("Vigencia Licencia:", dp.get("licencia_vigencia", "N/A"))
            ])
        
        if dp.get("antecedentes_legales"):
            datos_tabla.append(("Antecedentes Legales:", dp.get("antecedentes_legales_detalle", "Sí")))
        
        contactos = dp.get("contactos_emergencia", [])
        if contactos and len(contactos) > 0:
            contacto = contactos[0]
            datos_tabla.append(("Contacto de Emergencia:", 
                f"{contacto.get('nombre', 'N/A')} - {contacto.get('telefono', 'N/A')} ({contacto.get('relacion', 'N/A')})"))
        
        self._agregar_tabla_datos(doc, datos_tabla)
    
    def _agregar_salud(self, doc: Document, datos: Dict):
        """Agrega la sección de salud e intereses."""
        self._agregar_seccion(doc, "SALUD E INTERESES")
        
        salud = datos.get("salud_intereses", {})
        
        datos_salud = [
            ("Estado de Salud General:", salud.get("estado_salud", "N/A")),
            ("Tipo de Sangre:", salud.get("tipo_sangre", "N/A")),
            ("Alergias:", salud.get("alergias", "Ninguna")),
            ("Fuma:", "Sí" if salud.get("fuma") else "No"),
            ("Consume Alcohol:", "Sí" if salud.get("consume_alcohol") else "No")
        ]
        
        if salud.get("fuma"):
            datos_salud.append(("Frecuencia Tabaco:", salud.get("frecuencia_tabaco", "N/A")))
        if salud.get("consume_alcohol"):
            datos_salud.append(("Frecuencia Alcohol:", salud.get("frecuencia_alcohol", "N/A")))
        if salud.get("consume_otras_sustancias"):
            datos_salud.append(("Otras Sustancias:", salud.get("otras_sustancias_detalle", "Sí")))
        
        self._agregar_tabla_datos(doc, datos_salud)
        
        # Enfermedades crónicas
        enfermedades = salud.get("enfermedades_cronicas", [])
        if enfermedades and len(enfermedades) > 0:
            doc.add_paragraph("Enfermedades Crónicas:", style='Heading 3')
            
            table = doc.add_table(rows=len(enfermedades) + 1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            headers = ["Enfermedad", "Tratamiento", "Frecuencia Médico"]
            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            for idx, enf in enumerate(enfermedades, 1):
                row = table.rows[idx]
                row.cells[0].text = enf.get("nombre", "")
                row.cells[1].text = enf.get("tratamiento", "N/A")
                row.cells[2].text = enf.get("frecuencia_visita_medico", "N/A")
            
            doc.add_paragraph()
        
        # Actividades e intereses
        if salud.get("actividades_tiempo_libre"):
            p = doc.add_paragraph()
            p.add_run("Actividades de Tiempo Libre: ").bold = True
            p.add_run(salud.get("actividades_tiempo_libre", ""))
        
        if salud.get("deportes_practica"):
            p = doc.add_paragraph()
            p.add_run("Deportes que Practica: ").bold = True
            p.add_run(salud.get("deportes_practica", ""))
        
        doc.add_paragraph()
    
    def _agregar_informacion_familiar(self, doc: Document, datos: Dict):
        """Agrega la sección de información familiar."""
        self._agregar_seccion(doc, "INFORMACIÓN FAMILIAR")
        
        fam = datos.get("informacion_familiar", {})
        
        datos_generales = [
            ("Número de Hijos:", fam.get("numero_hijos", 0)),
            ("Ingreso Familiar Total:", f"${fam.get('ingreso_familiar_total', 0):,.2f}")
        ]
        
        self._agregar_tabla_datos(doc, datos_generales)
        
        # Miembros del hogar
        miembros = fam.get("miembros_hogar", [])
        if miembros:
            doc.add_paragraph("Miembros del Hogar:", style='Heading 3')
            
            table = doc.add_table(rows=len(miembros) + 1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            # Encabezados
            headers = ["Nombre", "Edad", "Parentesco", "Ocupación", "Ingreso"]
            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Datos
            for idx, miembro in enumerate(miembros, 1):
                row = table.rows[idx]
                row.cells[0].text = miembro.get("nombre", "")
                row.cells[1].text = str(miembro.get("edad", ""))
                row.cells[2].text = miembro.get("parentesco", "")
                row.cells[3].text = miembro.get("ocupacion", "")
                row.cells[4].text = f"${miembro.get('ingreso', 0):,.2f}"
            
            doc.add_paragraph()
    
    def _agregar_situacion_financiera(self, doc: Document, datos: Dict):
        """Agrega la sección de situación financiera."""
        self._agregar_seccion(doc, "SITUACIÓN FINANCIERA")
        
        fin = datos.get("situacion_financiera", {})
        
        # Información laboral
        datos_laborales = [
            ("Trabaja Actualmente:", "Sí" if fin.get("trabaja_actualmente", False) else "No"),
            ("Empresa:", fin.get("empresa_actual", "N/A")),
            ("Puesto:", fin.get("puesto_actual", "N/A")),
            ("Sueldo Mensual:", f"${fin.get('sueldo_mensual', 0):,.2f}"),
            ("Horario:", fin.get("horario", "N/A"))
        ]
        
        self._agregar_tabla_datos(doc, datos_laborales)
        
        # Gastos mensuales
        doc.add_paragraph("Gastos Mensuales:", style='Heading 3')
        
        gastos = fin.get("gastos", {})
        gastos_data = [
            ("Alimentación", f"${gastos.get('alimentacion', 0):,.2f}"),
            ("Salud", f"${gastos.get('salud', 0):,.2f}"),
            ("Educación", f"${gastos.get('educacion', 0):,.2f}"),
            ("Vivienda", f"${gastos.get('vivienda', 0):,.2f}"),
            ("Transporte", f"${gastos.get('transporte', 0):,.2f}"),
            ("Servicios", f"${gastos.get('servicios', 0):,.2f}"),
            ("Otros", f"${gastos.get('otros', 0):,.2f}"),
            ("TOTAL", f"${gastos.get('total', 0):,.2f}")
        ]
        
        self._agregar_tabla_datos(doc, gastos_data)
        
        # Balance
        balance = fin.get("balance", 0)
        p = doc.add_paragraph()
        p.add_run("Balance Mensual (Ingreso - Gastos): ").bold = True
        balance_run = p.add_run(f"${balance:,.2f}")
        balance_run.font.bold = True
        balance_run.font.color.rgb = RGBColor(0, 128, 0) if balance >= 0 else RGBColor(255, 0, 0)
        
        # Observaciones
        if fin.get("observaciones_financieras"):
            doc.add_paragraph()
            doc.add_paragraph("Observaciones:", style='Heading 3')
            doc.add_paragraph(fin.get("observaciones_financieras", ""))
        
        doc.add_paragraph()
    
    def _agregar_empleo_actual(self, doc: Document, datos: Dict):
        """Agrega la sección de empleo actual detallado."""
        empleo = datos.get("empleo_actual", {})
        
        if not empleo.get("tiene_empleo"):
            return
        
        self._agregar_seccion(doc, "EMPLEO ACTUAL")
        
        datos_empleo = [
            ("Empresa:", empleo.get("empresa", "N/A")),
            ("Puesto:", empleo.get("puesto", "N/A")),
            ("Área/Departamento:", empleo.get("area_departamento", "N/A")),
            ("Antigüedad:", empleo.get("antiguedad", "N/A")),
            ("Tipo de Contrato:", empleo.get("tipo_contrato", "N/A")),
            ("Teléfono de la Empresa:", empleo.get("telefono_empresa", "N/A")),
            ("Dirección de la Empresa:", empleo.get("direccion_empresa", "N/A")),
            ("Jefe Directo:", empleo.get("nombre_jefe", "N/A")),
            ("Puesto del Jefe:", empleo.get("puesto_jefe", "N/A"))
        ]
        
        self._agregar_tabla_datos(doc, datos_empleo)
    
    def _agregar_estilo_vida(self, doc: Document, datos: Dict):
        """Agrega la sección de estilo de vida."""
        estilo = datos.get("estilo_vida", {})
        
        if not estilo:
            return
        
        self._agregar_seccion(doc, "ESTILO DE VIDA")
        
        datos_estilo = []
        
        if estilo.get("vehiculo_propio"):
            datos_estilo.append(("Vehículo Propio:", 
                f"{estilo.get('marca_vehiculo', '')} {estilo.get('modelo_vehiculo', '')} ({estilo.get('ano_vehiculo', 'N/A')})"))
        else:
            datos_estilo.append(("Vehículo Propio:", "No"))
        
        if estilo.get("viajes_ultimo_ano"):
            datos_estilo.append(("Viajes en el Último Año:", estilo.get("destinos_viajes", "Sí")))
        
        if estilo.get("hobbies"):
            datos_estilo.append(("Hobbies:", estilo.get("hobbies", "")))
        
        if estilo.get("pertenece_asociaciones"):
            datos_estilo.append(("Asociaciones/Clubes:", estilo.get("asociaciones_detalle", "Sí")))
        
        if datos_estilo:
            self._agregar_tabla_datos(doc, datos_estilo)
    
    def _agregar_vivienda(self, doc: Document, datos: Dict):
        """Agrega la sección de vivienda."""
        self._agregar_seccion(doc, "VIVIENDA Y PATRIMONIO")
        
        viv = datos.get("vivienda", {})
        
        datos_vivienda = [
            ("Tipo de Vivienda:", viv.get("tipo_vivienda", "N/A")),
            ("Tenencia:", viv.get("tenencia", "N/A")),
            ("Zona:", viv.get("tipo_zona", "N/A")),
            ("Materiales:", viv.get("materiales_construccion", "N/A")),
            ("Tiempo de Residencia:", viv.get("tiempo_residencia", "N/A")),
            ("Número de Cuartos:", viv.get("numero_cuartos", "N/A"))
        ]
        
        self._agregar_tabla_datos(doc, datos_vivienda)
        
        # Servicios
        servicios = viv.get("servicios", {})
        servicios_lista = [k.replace("_", " ").title() for k, v in servicios.items() if v]
        
        if servicios_lista:
            p = doc.add_paragraph()
            p.add_run("Servicios: ").bold = True
            p.add_run(", ".join(servicios_lista))
        
        doc.add_paragraph()
    
    def _agregar_historial_laboral(self, doc: Document, datos: Dict):
        """Agrega la sección de historial laboral."""
        historial = datos.get("historial_laboral", [])
        
        if not historial:
            return
        
        self._agregar_seccion(doc, "HISTORIAL LABORAL")
        
        table = doc.add_table(rows=len(historial) + 1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Encabezados
        headers = ["Empresa", "Puesto", "Periodo", "Motivo de Salida"]
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Datos
        for idx, empleo in enumerate(historial, 1):
            row = table.rows[idx]
            row.cells[0].text = empleo.get("empresa", "")
            row.cells[1].text = empleo.get("puesto", "")
            periodo = f"{empleo.get('fecha_inicio', '')} - {empleo.get('fecha_fin', '')}"
            row.cells[2].text = periodo
            row.cells[3].text = empleo.get("motivo_separacion", "")
        
        doc.add_paragraph()
    
    def _agregar_referencias(self, doc: Document, datos: Dict):
        """Agrega la sección de referencias."""
        referencias = datos.get("referencias", [])
        
        if not referencias:
            return
        
        self._agregar_seccion(doc, "REFERENCIAS PERSONALES")
        
        table = doc.add_table(rows=len(referencias) + 1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        # Encabezados
        headers = ["Nombre", "Relación", "Teléfono", "Tiempo Conocido"]
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Datos
        for idx, ref in enumerate(referencias, 1):
            row = table.rows[idx]
            row.cells[0].text = ref.get("nombre", "")
            row.cells[1].text = ref.get("relacion", "")
            row.cells[2].text = ref.get("telefono", "")
            row.cells[3].text = ref.get("tiempo_conocido", "")
        
        doc.add_paragraph()
    
    def _agregar_analisis_riesgos(self, doc: Document, datos: Dict):
        """Agrega la sección de análisis de riesgos con justificaciones."""
        from src.logic.calculador_riesgos import CalculadorRiesgos
        
        self._agregar_seccion(doc, "ANÁLISIS DE RIESGOS")
        
        # Crear instancia del calculador
        calc = CalculadorRiesgos(datos)
        
        # Calcular todos los riesgos con justificaciones
        resultados = calc.calcular_todos_riesgos()
        
        # Tabla resumen de riesgos
        riesgos_data = []
        
        categorias = [
            ("financiero", "Riesgo Financiero"),
            ("familiar", "Riesgo Familiar"),
            ("vivienda", "Riesgo Vivienda"),
            ("laboral", "Riesgo Laboral"),
            ("salud", "Riesgo Salud"),
            ("estilo_vida", "Riesgo Estilo de Vida")
        ]
        
        for key, label in categorias:
            if key in resultados:
                nivel = resultados[key]["puntaje"]
                interpretacion = CalculadorRiesgos.obtener_interpretacion_riesgo(nivel)
                riesgos_data.append((label, f"{nivel} - {interpretacion}"))
        
        # Riesgo global
        if "global" in resultados:
            nivel = resultados["global"]["puntaje"]
            interpretacion = CalculadorRiesgos.obtener_interpretacion_riesgo(nivel)
            riesgos_data.append(("RIESGO GLOBAL", f"{nivel} - {interpretacion}"))
        
        self._agregar_tabla_datos(doc, riesgos_data)
        
        # JUSTIFICACIONES DETALLADAS
        doc.add_heading("Justificaciones y Detalles", level=3)
        
        for key, label in categorias:
            if key in resultados and resultados[key]["justificaciones"]:
                # Título de la categoría
                p = doc.add_paragraph()
                p.add_run(f"{label}:").bold = True
                
                # Lista de justificaciones
                justificaciones = resultados[key]["justificaciones"]
                for just in justificaciones:
                    doc.add_paragraph(just, style='List Bullet')
                
                doc.add_paragraph()  # Espaciado entre categorías
    
    def _agregar_conclusiones(self, doc: Document, datos: Dict):
        """Agrega la sección de conclusiones."""
        conclusiones = datos.get("conclusiones", "")
        
        if not conclusiones:
            return
        
        self._agregar_seccion(doc, "CONCLUSIONES Y RECOMENDACIONES")
        
        # Agregar conclusiones en un párrafo con fondo
        p = doc.add_paragraph(conclusiones)
        p.style = 'Intense Quote'
        
        doc.add_paragraph()
    
    def _agregar_fotografias(self, doc: Document, datos: Dict):
        """Agrega la sección de fotografías."""
        fotos = datos.get("fotos", [])
        
        if not fotos:
            return
        
        doc.add_page_break()
        self._agregar_seccion(doc, "EVIDENCIA FOTOGRÁFICA")
        
        for foto in fotos:
            archivo = foto.get("archivo", "")
            tipo = foto.get("tipo", "Sin categoría")
            descripcion = foto.get("descripcion", "")
            
            if archivo and os.path.exists(archivo):
                try:
                    # Agregar la imagen
                    paragraph = doc.add_paragraph()
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(archivo, width=Inches(4))
                    
                    # Pie de foto
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption.add_run(f"{tipo}")
                    caption_run.font.bold = True
                    
                    if descripcion:
                        caption.add_run(f" - {descripcion}")
                    
                    doc.add_paragraph()
                    
                except Exception as e:
                    print(f"Error al incluir fotografía: {e}")
    
    def exportar(self, estudio_datos: Dict, ruta_salida: str) -> bool:
        """
        Exporta un estudio socioeconómico a formato Word.
        
        Args:
            estudio_datos: Diccionario con los datos del estudio.
            ruta_salida: Ruta completa del archivo DOCX de salida.
            
        Returns:
            True si se exportó correctamente, False en caso contrario.
        """
        try:
            os.makedirs(os.path.dirname(ruta_salida), exist_ok=True)
            
            doc = Document()
            
            # Configurar márgenes
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Crear documento
            self._agregar_encabezado(doc)
            
            # Título del reporte
            nombre = estudio_datos.get("datos_personales", {}).get("nombre_completo", "Sin nombre")
            title = doc.add_heading(f"ESTUDIO SOCIOECONÓMICO", level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtitle = doc.add_heading(nombre, level=2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            fecha_p = doc.add_paragraph()
            fecha_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fecha_p.add_run(f"Fecha de elaboración: {datetime.now().strftime('%d/%m/%Y')}")
            
            doc.add_paragraph()
            
            # Secciones del reporte
            self._agregar_datos_personales(doc, estudio_datos)
            self._agregar_salud(doc, estudio_datos)
            self._agregar_informacion_familiar(doc, estudio_datos)
            self._agregar_situacion_financiera(doc, estudio_datos)
            self._agregar_empleo_actual(doc, estudio_datos)
            self._agregar_estilo_vida(doc, estudio_datos)
            self._agregar_vivienda(doc, estudio_datos)
            self._agregar_historial_laboral(doc, estudio_datos)
            self._agregar_referencias(doc, estudio_datos)
            self._agregar_analisis_riesgos(doc, estudio_datos)
            self._agregar_conclusiones(doc, estudio_datos)
            self._agregar_fotografias(doc, estudio_datos)
            
            # Guardar documento
            doc.save(ruta_salida)
            
            return True
            
        except Exception as e:
            print(f"Error al exportar Word: {e}")
            import traceback
            traceback.print_exc()
            return False
